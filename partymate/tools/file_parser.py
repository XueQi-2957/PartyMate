"""
PartyMate 文件解析引擎

支持格式：
  - PDF (*.pdf)       → pymupdf 提取文本
  - Word (*.docx)     → python-docx 提取文本
  - 图片 (*.png/jpg)  → easyocr 光学字符识别

输出统一为:
  {
    "filename": "xxx.pdf",
    "type": "pdf" | "docx" | "image",
    "text": "全文内容",
    "pages": 3,
    "preview": "前500字预览",
    "error": None | str
  }
"""

from __future__ import annotations

import io
import os
import re
from pathlib import Path

import fitz  # pymupdf

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# lazy-load easyocr (large model, only when needed)
_OCR_READER = None


def _get_ocr_reader():
    global _OCR_READER
    if _OCR_READER is None:
        import easyocr
        _OCR_READER = easyocr.Reader(["ch_sim", "en"], gpu=False)
    return _OCR_READER


SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


def parse_file(filepath: str | Path) -> dict:
    """
    解析文件，返回统一结构
    """
    path = Path(filepath)
    if not path.exists():
        return {
            "filename": path.name,
            "type": "unknown",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": f"文件不存在: {path}",
        }

    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext in (".docx", ".doc"):
        return _parse_docx(path)
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff"):
        return _parse_image(path)
    else:
        return {
            "filename": path.name,
            "type": "unknown",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": f"不支持的文件格式: {ext}",
        }


def _parse_pdf(path: Path) -> dict:
    """解析 PDF 文件"""
    try:
        doc = fitz.open(path)
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                pages.append(text)
        full_text = "\n\n".join(pages)
        doc.close()
        return {
            "filename": path.name,
            "type": "pdf",
            "text": full_text,
            "pages": len(pages),
            "preview": _preview(full_text),
            "ocr_segments": [],
            "error": None,
        }
    except Exception as e:
        return {
            "filename": path.name,
            "type": "pdf",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": f"PDF 解析失败: {e}",
        }


def _parse_docx(path: Path) -> dict:
    """解析 Word 文件"""
    if not DOCX_OK:
        return {
            "filename": path.name,
            "type": "docx",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": "python-docx 未安装",
        }
    try:
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables.append(" | ".join(cells))
        full_text = "\n".join(paragraphs)
        if tables:
            full_text += "\n\n【表格内容】\n" + "\n".join(tables)
        return {
            "filename": path.name,
            "type": "docx",
            "text": full_text,
            "pages": 1,
            "preview": _preview(full_text),
            "ocr_segments": [],
            "error": None,
        }
    except Exception as e:
        return {
            "filename": path.name,
            "type": "docx",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": f"Word 解析失败: {e}",
        }


def _parse_image(path: Path) -> dict:
    """解析图片（OCR）"""
    try:
        reader = _get_ocr_reader()
        results = reader.readtext(str(path))
        lines = []
        segments = []
        for (bbox, text, confidence) in results:
            if confidence > 0.3:
                lines.append(text)
                segments.append(
                    {
                        "text": text,
                        "confidence": confidence,
                        "bbox": bbox,
                    }
                )
        full_text = "\n".join(lines)
        return {
            "filename": path.name,
            "type": "image",
            "text": full_text,
            "pages": 1,
            "preview": _preview(full_text),
            "ocr_segments": segments,
            "error": None,
        }
    except ImportError:
        return {
            "filename": path.name,
            "type": "image",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": "easyocr 未安装",
        }
    except Exception as e:
        return {
            "filename": path.name,
            "type": "image",
            "text": "",
            "pages": 0,
            "preview": "",
            "ocr_segments": [],
            "error": f"图片 OCR 失败: {e}",
        }


def _preview(text: str, max_chars: int = 500) -> str:
    """生成文本预览"""
    text = text.strip()
    if not text:
        return "(空内容)"
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "...\n\n[预览截断，共 {} 字]".format(len(text))


def preview_from_bytes(filename: str, data: bytes) -> dict:
    """从字节流解析文件（用于 Web 上传）"""
    ext = Path(filename).suffix.lower()
    # 写入临时文件
    tmp_dir = Path(os.getcwd()) / ".tmp_uploads"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = tmp_dir / filename
    tmp_path.write_bytes(data)
    try:
        result = parse_file(tmp_path)
        return result
    finally:
        # 只保留解析结果，清理临时文件
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass
