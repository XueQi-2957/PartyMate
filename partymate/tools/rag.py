"""
PartyMate hybrid RAG engine.

The index is built as parent/child retrieval:
  - parent chunks keep document, chapter, table, and template context.
  - child chunks are the searchable evidence units stored in ChromaDB and BM25.
  - Ollama bge-m3 creates dense vectors.
  - Ollama minicpm-v-fast may mark semantic boundaries, but original text is
    always sliced from source lines so the model cannot rewrite evidence.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import sys
import time
import uuid
from dataclasses import asdict, dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

import chromadb
import httpx
from rank_bm25 import BM25Okapi

# ── Paths ─────────────────────────────────────────────

HERE = Path(__file__).resolve().parent.parent
RUNTIME_KNOWLEDGE_DIR = HERE / "knowledge"
SOURCE_KNOWLEDGE_DIR = HERE.parent / "知识库"
CHROMA_DIR = RUNTIME_KNOWLEDGE_DIR / "chroma_db"
CHUNK_INDEX_FILE = RUNTIME_KNOWLEDGE_DIR / "chunks_index.json"
PARENT_INDEX_FILE = RUNTIME_KNOWLEDGE_DIR / "parents_index.json"

# ── Ollama / OpenAI-compatible config ─────────────────

API_BASE = os.getenv("PARTYMATE_API_BASE") or os.getenv("OPENAI_BASE_URL") or "http://127.0.0.1:11434/v1"
API_KEY = os.getenv("PARTYMATE_API_KEY") or os.getenv("OPENAI_API_KEY") or ""
# Runtime answer/rerank model remains user-configurable. Semantic chunking has
# its own fast local default below.
MODEL = os.getenv("PARTYMATE_MODEL") or os.getenv("HERMES_MODEL") or "gemma4:e4b"
CHUNK_MODEL = os.getenv("PARTYMATE_CHUNK_MODEL") or "minicpm-v-fast:latest"
VISION_MODEL = os.getenv("PARTYMATE_VISION_MODEL") or "gemma4:e4b"
USE_LLM_CHUNKING = os.getenv("PARTYMATE_USE_LLM_CHUNKING", "0").lower() in {"1", "true", "yes"}

OLLAMA_BASE = os.getenv("OLLAMA_HOST") or "http://127.0.0.1:11434"
EMBED_MODEL = os.getenv("PARTYMATE_EMBED_MODEL") or "bge-m3:latest"

# ── Retrieval parameters ──────────────────────────────

DEFAULT_TOP_K = 5
HYBRID_ALPHA = 0.6
HYBRID_TOP_K = 15
DEFAULT_RERANK = True
RERANK_MODE = os.getenv("PARTYMATE_RERANK_MODE") or "cross-encoder"
RERANK_MODEL_NAME = os.getenv("PARTYMATE_RERANK_MODEL") or "BAAI/bge-reranker-v2-m3"
FUSION_MODE = os.getenv("PARTYMATE_FUSION_MODE") or "rrf"  # "linear" 或 "rrf"
RRF_K = 60  # RRF 常数

PARENT_TARGET_CHARS = 2200
PARENT_MAX_CHARS = 3200
CHILD_MIN_CHARS = 12
CHILD_TARGET_CHARS = 450
CHILD_MAX_CHARS = 700
TABLE_CHILD_CHARS = 1400

SUPPORTED_EXTS = {".docx", ".pdf", ".txt"}
SKIPPED_EXTS = {".doc"}


@dataclass
class DocumentElement:
    doc_name: str
    source_path: str
    page_no: int | None
    heading_path: str
    element_type: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ParentChunk:
    parent_id: str
    doc_name: str
    source_path: str
    title: str
    heading_path: str
    content: str
    chunk_type: str = "text"
    page_no: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "ParentChunk":
        return cls(**data)


@dataclass
class ChildChunk:
    child_id: str
    parent_id: str
    doc_name: str
    source_path: str
    title: str
    heading_path: str
    content: str
    chunk_type: str = "text"
    page_no: int | None = None
    article_no: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def citation(self) -> str:
        parts = []
        if self.doc_name:
            parts.append(self.doc_name)
        if self.heading_path and not self.heading_path.startswith(self.doc_name):
            parts.append(self.heading_path)
        elif self.heading_path:
            parts = [self.heading_path]
        if self.article_no and self.article_no not in self.heading_path:
            parts.append(self.article_no)
        if self.page_no:
            parts.append(f"第{self.page_no}页")
        return " > ".join([p for p in parts if p])

    @property
    def embedding_text(self) -> str:
        prefix = self.citation
        if self.chunk_type == "table":
            prefix = f"{prefix} > 表格"
        return f"{prefix}\n{self.content}".strip()

    def to_metadata(self, idx: int, build_at: str) -> dict[str, str]:
        meta = {
            "title": self.title,
            "idx": str(idx),
            "build_at": build_at,
            "parent_id": self.parent_id,
            "doc_name": self.doc_name,
            "source_path": self.source_path,
            "heading_path": self.heading_path,
            "article_no": self.article_no,
            "chunk_type": self.chunk_type,
            "citation": self.citation,
            "page_no": str(self.page_no or ""),
        }
        for key, value in self.metadata.items():
            if isinstance(value, (dict, list)):
                meta[key] = json.dumps(value, ensure_ascii=False)
            else:
                meta[key] = "" if value is None else str(value)
        return meta


def _clean_text(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.strip() for line in text.splitlines()).strip()


def _safe_id(prefix: str = "") -> str:
    raw = uuid.uuid4().hex[:12]
    return f"{prefix}{raw}" if prefix else raw


def _looks_like_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) > 80:
        return False
    patterns = [
        r"^第[一二三四五六七八九十百\d]+[章节条阶段]\s*",
        r"^[一二三四五六七八九十\d]+[、．.]\s*",
        r"^（[一二三四五六七八九十\d]+）",
        r".*(条例|细则|规程|办法|规定|标准|手册|说明|清单)$",
    ]
    return any(re.match(pattern, stripped) for pattern in patterns)


def _article_no(text: str) -> str:
    match = re.match(r"^\s*(第[一二三四五六七八九十百零〇两\d]+条)", text)
    if match:
        return match.group(1)
    step_match = re.match(r"^\s*(第[一二三四五六七八九十百零〇两\d]+步)", text)
    if step_match:
        return step_match.group(1)
    return ""


def _discover_knowledge_files(path: str | Path | None = None) -> tuple[list[Path], list[Path]]:
    root = Path(path) if path is not None else SOURCE_KNOWLEDGE_DIR
    if root.is_file():
        files = [root]
    elif root.exists():
        files = sorted(p for p in root.iterdir() if p.is_file())
    else:
        return [], []
    supported = [p for p in files if p.suffix.lower() in SUPPORTED_EXTS]
    skipped = [p for p in files if p.suffix.lower() in SKIPPED_EXTS]
    return supported, skipped


def _load_txt_elements(path: str | Path) -> list[DocumentElement]:
    source = Path(path)
    text = source.read_text(encoding="utf-8", errors="ignore")
    lines = [_clean_text(line) for line in text.splitlines()]
    elements: list[DocumentElement] = []
    heading_parts: list[str] = []
    for line in lines:
        if not line:
            continue
        if _looks_like_heading(line):
            if not heading_parts:
                heading_parts = [line]
            elif re.match(r"^第[一二三四五六七八九十百\d]+[章节阶段]", line):
                heading_parts = [heading_parts[0], line]
            elif len(heading_parts) < 3:
                heading_parts.append(line)
            else:
                heading_parts[-1] = line
            element_type = "heading"
        else:
            element_type = "text"
        elements.append(
            DocumentElement(
                doc_name=source.stem,
                source_path=str(source),
                page_no=None,
                heading_path=" > ".join(heading_parts) if heading_parts else source.stem,
                element_type=element_type,
                text=line,
            )
        )
    return elements


def _iter_docx_body(doc: Any) -> list[tuple[str, Any]]:
    from docx.oxml.ns import qn

    body = doc.element.body
    out: list[tuple[str, Any]] = []
    para_idx = 0
    table_idx = 0
    for child in body:
        if child.tag == qn("w:p"):
            out.append(("paragraph", doc.paragraphs[para_idx]))
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            out.append(("table", doc.tables[table_idx]))
            table_idx += 1
    return out


def _dedupe_row(row: list[str]) -> list[str]:
    result: list[str] = []
    last = object()
    for cell in row:
        if cell != last:
            result.append(cell)
        last = cell
    return result


def _load_docx_elements(path: str | Path) -> list[DocumentElement]:
    from docx import Document

    source = Path(path)
    doc = Document(str(source))
    elements: list[DocumentElement] = []
    heading_parts: list[str] = []
    table_num = 0

    for kind, obj in _iter_docx_body(doc):
        if kind == "paragraph":
            text = _clean_text(obj.text)
            if not text:
                continue
            style_name = obj.style.name if obj.style else ""
            is_heading = "heading" in style_name.lower() or _looks_like_heading(text)
            if is_heading:
                if not heading_parts:
                    heading_parts = [text]
                elif re.match(r"^第[一二三四五六七八九十百\d]+[章节阶段]", text):
                    heading_parts = [heading_parts[0], text]
                elif len(heading_parts) < 3:
                    heading_parts.append(text)
                else:
                    heading_parts[-1] = text
            elements.append(
                DocumentElement(
                    doc_name=source.stem,
                    source_path=str(source),
                    page_no=None,
                    heading_path=" > ".join(heading_parts) if heading_parts else source.stem,
                    element_type="heading" if is_heading else "text",
                    text=text,
                )
            )
            continue

        table_num += 1
        rows = []
        for row in obj.rows:
            cells = [_clean_text(cell.text) for cell in row.cells]
            rows.append(_dedupe_row(cells))
        if not rows:
            continue
        headers = rows[0]
        body_rows = rows[1:] if len(rows) > 1 else []
        table_title = heading_parts[-1] if heading_parts else f"表格{table_num}"
        text = _serialize_table_rows(table_title, headers, body_rows)
        elements.append(
            DocumentElement(
                doc_name=source.stem,
                source_path=str(source),
                page_no=None,
                heading_path=" > ".join(heading_parts) if heading_parts else source.stem,
                element_type="table",
                text=text,
                metadata={
                    "table_num": table_num,
                    "headers": headers,
                    "rows": body_rows,
                    "num_rows": len(body_rows),
                    "num_cols": len(headers),
                },
            )
        )
    return elements


def _load_pdf_elements(path: str | Path) -> list[DocumentElement]:
    import fitz

    source = Path(path)
    elements: list[DocumentElement] = []
    with fitz.open(str(source)) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = _clean_text(page.get_text("text"))
            if not text:
                continue
            elements.append(
                DocumentElement(
                    doc_name=source.stem,
                    source_path=str(source),
                    page_no=page_index,
                    heading_path=f"{source.stem} > 第{page_index}页",
                    element_type="text",
                    text=text,
                )
            )
    return elements


def _serialize_table_rows(title: str, headers: list[str], rows: list[list[str]]) -> str:
    if not rows:
        return f"表格：{title}\n字段：" + "、".join(h for h in headers if h)

    lines = [f"表格：{title}", "字段：" + "、".join(h for h in headers if h)]
    for row_idx, row in enumerate(rows, 1):
        pairs: list[str] = []
        for col_idx, cell in enumerate(row):
            if not cell:
                continue
            header = headers[col_idx] if col_idx < len(headers) and headers[col_idx] else f"字段{col_idx + 1}"
            pairs.append(f"{header}: {cell}")
        if pairs:
            lines.append(f"第{row_idx}行: " + "；".join(pairs))
    return "\n".join(lines)


def _load_document_elements(path: str | Path) -> list[DocumentElement]:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".docx":
        return _load_docx_elements(source)
    if suffix == ".pdf":
        return _load_pdf_elements(source)
    if suffix == ".txt":
        return _load_txt_elements(source)
    return []


def _elements_to_parents(elements: list[DocumentElement]) -> list[ParentChunk]:
    parents: list[ParentChunk] = []
    buffer: list[DocumentElement] = []
    current_key = ""

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        first = buffer[0]
        content = _clean_text("\n".join(e.text for e in buffer if e.text))
        if not content:
            buffer = []
            return
        title = first.heading_path or first.doc_name
        chunk_type = "table" if all(e.element_type == "table" for e in buffer) else "text"
        metadata: dict[str, Any] = {}
        if chunk_type == "table":
            metadata = dict(first.metadata)
        parents.append(
            ParentChunk(
                parent_id=_safe_id("p_"),
                doc_name=first.doc_name,
                source_path=first.source_path,
                title=title,
                heading_path=first.heading_path or first.doc_name,
                content=content,
                chunk_type=chunk_type,
                page_no=first.page_no,
                metadata=metadata,
            )
        )
        buffer = []

    for element in elements:
        key = f"{element.source_path}|{element.heading_path}|{element.page_no}|{element.element_type}"
        should_flush = False
        if not buffer:
            current_key = key
        elif element.element_type == "table" or buffer[-1].element_type == "table":
            should_flush = True
        elif key != current_key and sum(len(e.text) for e in buffer) >= CHILD_MIN_CHARS:
            should_flush = True
        elif sum(len(e.text) for e in buffer) + len(element.text) > PARENT_MAX_CHARS:
            should_flush = True

        if should_flush:
            flush()
            current_key = key
        buffer.append(element)

    flush()
    return parents


def _line_numbered_text(lines: list[str]) -> str:
    return "\n".join(f"{i + 1}: {line}" for i, line in enumerate(lines))


def _extract_json_array(text: str) -> list[dict[str, Any]]:
    if not text.strip():
        return []
    fenced = re.search(r"```(?:json)?\s*(\[.*?\])\s*```", text, flags=re.S)
    raw = fenced.group(1) if fenced else text
    if not raw.strip().startswith("["):
        start = raw.find("[")
        end = raw.rfind("]")
        if start >= 0 and end > start:
            raw = raw[start : end + 1]
    try:
        data = json.loads(raw)
    except Exception:
        return []
    return data if isinstance(data, list) else []


_BOUNDARY_SYSTEM_PROMPT = """你是党务规范文档的语义边界标注器。
你只能返回 JSON 数组，不能改写、总结、补充原文。
每个对象必须包含 title、start_line、end_line、chunk_type、article_no、reason。
start_line 和 end_line 必须来自用户提供的行号，且包含完整句子或完整条款。
正文 chunk 目标 300-500 个中文字符；很短但完整的条款可以单独成块。

Few-shot:
输入:
1: 第一章 总则
2: 第一条 发展党员工作应当贯彻党的基本理论。
3: 第二条 党组织应当把政治标准放在首位。
输出:
[
  {"title":"第一章 总则 > 第一条","start_line":2,"end_line":2,"chunk_type":"text","article_no":"第一条","reason":"完整条款"},
  {"title":"第一章 总则 > 第二条","start_line":3,"end_line":3,"chunk_type":"text","article_no":"第二条","reason":"完整条款"}
]
"""


def _call_llm(system: str, prompt: str, max_tokens: int = 2048, model: str | None = None) -> str:
    headers = {"Content-Type": "application/json"}
    if API_KEY:
        headers["Authorization"] = f"Bearer {API_KEY}"
    body = {
        "model": model or MODEL,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.0,
        "max_tokens": max_tokens,
    }
    try:
        with httpx.Client(timeout=90.0) as client:
            resp = client.post(f"{API_BASE}/chat/completions", headers=headers, json=body)
        if resp.status_code == 200:
            return resp.json()["choices"][0]["message"]["content"]
    except Exception:
        return ""
    return ""


def _llm_child_boundaries(parent: ParentChunk) -> list[dict[str, Any]]:
    lines = [line for line in parent.content.splitlines() if line.strip()]
    if len(lines) < 2 or len(parent.content) > 7000:
        return []
    prompt = (
        f"文档标题路径：{parent.heading_path}\n"
        f"块类型：{parent.chunk_type}\n"
        "请按语义边界标注以下文本：\n\n"
        f"{_line_numbered_text(lines)}"
    )
    result = _call_llm(_BOUNDARY_SYSTEM_PROMPT, prompt, model=CHUNK_MODEL)
    boundaries = _extract_json_array(result)
    valid: list[dict[str, Any]] = []
    for item in boundaries:
        try:
            start = int(item["start_line"])
            end = int(item["end_line"])
        except Exception:
            continue
        if start < 1 or end < start or end > len(lines):
            continue
        valid.append(
            {
                "title": str(item.get("title") or parent.title),
                "start_line": start,
                "end_line": end,
                "chunk_type": str(item.get("chunk_type") or parent.chunk_type),
                "article_no": str(item.get("article_no") or ""),
            }
        )
    return valid


def _split_long_text(lines: list[str], max_chars: int = CHILD_MAX_CHARS) -> list[tuple[int, int]]:
    spans: list[tuple[int, int]] = []
    start = 0
    count = 0
    for idx, line in enumerate(lines):
        count += len(line)
        is_boundary = bool(re.search(r"[。！？；]$", line.strip())) or len(line) < 80
        if count >= max_chars and is_boundary:
            spans.append((start, idx))
            start = idx + 1
            count = 0
    if start < len(lines):
        spans.append((start, len(lines) - 1))
    return spans


def _rule_child_boundaries(parent: ParentChunk) -> list[dict[str, Any]]:
    lines = [line for line in parent.content.splitlines() if line.strip()]
    if not lines:
        return []

    if parent.chunk_type == "table":
        spans: list[tuple[int, int]] = []
        start = 0
        chars = 0
        for idx, line in enumerate(lines):
            chars += len(line)
            if idx > 1 and line.startswith("第") and "行:" in line and chars > TABLE_CHILD_CHARS:
                spans.append((start, idx - 1))
                start = idx
                chars = len(line)
        spans.append((start, len(lines) - 1))
        return [
            {
                "title": parent.title,
                "start_line": s + 1,
                "end_line": e + 1,
                "chunk_type": "table",
                "article_no": "",
            }
            for s, e in spans
        ]

    article_starts = [i for i, line in enumerate(lines) if _article_no(line)]
    if article_starts:
        spans = []
        for pos, start in enumerate(article_starts):
            end = article_starts[pos + 1] - 1 if pos + 1 < len(article_starts) else len(lines) - 1
            spans.extend(_split_long_text(lines[start : end + 1], max_chars=CHILD_MAX_CHARS))
            if spans and spans[-1][0] < start:
                # The long-text splitter works on sliced indexes.
                pass
        normalized: list[dict[str, Any]] = []
        for pos, start in enumerate(article_starts):
            end = article_starts[pos + 1] - 1 if pos + 1 < len(article_starts) else len(lines) - 1
            slice_spans = _split_long_text(lines[start : end + 1], max_chars=CHILD_MAX_CHARS)
            for rel_s, rel_e in slice_spans:
                abs_s = start + rel_s
                abs_e = start + rel_e
                article = _article_no(lines[start])
                normalized.append(
                    {
                        "title": f"{parent.heading_path} > {article}" if article else parent.title,
                        "start_line": abs_s + 1,
                        "end_line": abs_e + 1,
                        "chunk_type": "text",
                        "article_no": article,
                    }
                )
        return normalized

    spans = _split_long_text(lines, max_chars=CHILD_TARGET_CHARS)
    return [
        {
            "title": parent.title,
            "start_line": s + 1,
            "end_line": e + 1,
            "chunk_type": parent.chunk_type,
            "article_no": _article_no(lines[s]),
        }
        for s, e in spans
    ]


def _children_from_parent(parent: ParentChunk, use_llm: bool = True) -> list[ChildChunk]:
    lines = [line for line in parent.content.splitlines() if line.strip()]
    if not lines:
        return []
    if parent.chunk_type == "text" and len(lines) == 1 and _looks_like_heading(lines[0]) and not _article_no(lines[0]):
        return []
    boundaries = _llm_child_boundaries(parent) if use_llm and parent.chunk_type == "text" else []
    if not boundaries:
        boundaries = _rule_child_boundaries(parent)

    children: list[ChildChunk] = []
    seen_spans: set[tuple[int, int]] = set()
    for boundary in boundaries:
        start = int(boundary["start_line"]) - 1
        end = int(boundary["end_line"]) - 1
        if start < 0 or end < start or end >= len(lines):
            continue
        span = (start, end)
        if span in seen_spans:
            continue
        seen_spans.add(span)
        content = _clean_text("\n".join(lines[start : end + 1]))
        if len(content) < CHILD_MIN_CHARS and parent.chunk_type != "table":
            continue
        article = str(boundary.get("article_no") or _article_no(content))
        child_type = str(boundary.get("chunk_type") or parent.chunk_type)
        title = str(boundary.get("title") or parent.title)
        metadata = dict(parent.metadata)
        if child_type == "table":
            row_numbers = [
                int(m.group(1))
                for line in lines[start : end + 1]
                if (m := re.match(r"^第(\d+)行:", line.strip()))
            ]
            if row_numbers:
                metadata["row_range"] = f"{min(row_numbers)}-{max(row_numbers)}"
        children.append(
            ChildChunk(
                child_id=_safe_id("c_"),
                parent_id=parent.parent_id,
                doc_name=parent.doc_name,
                source_path=parent.source_path,
                title=title,
                heading_path=parent.heading_path,
                content=content,
                chunk_type=child_type,
                page_no=parent.page_no,
                article_no=article,
                metadata=metadata,
            )
        )
    return children


def _ollama_embed(texts: list[str]) -> list[list[float]]:
    if not texts:
        return []
    try:
        with httpx.Client(timeout=90.0) as client:
            resp = client.post(f"{OLLAMA_BASE}/api/embed", json={"model": EMBED_MODEL, "input": texts})
        if resp.status_code == 200:
            payload = resp.json()
            return payload.get("embeddings", [])
    except Exception:
        return []
    return []


class _OllamaEmbeddingFunction:
    def __call__(self, input: list[str]) -> list[list[float]]:
        return _ollama_embed(input)

    def name(self) -> str:
        return f"ollama_{EMBED_MODEL}"


def _tokenize(text: str) -> list[str]:
    chinese = re.findall(r"[\u4e00-\u9fff]{1,2}", text)
    words = re.findall(r"[A-Za-z0-9_]+", text.lower())
    return [t for t in chinese + words if t.strip()]


def extract_docx_text(docx_path: str | Path) -> str:
    elements = _load_docx_elements(docx_path)
    return "\n".join(e.text for e in elements if e.element_type != "table")


def _extract_interleaved(docx_path: str | Path) -> dict[str, Any]:
    elements = _load_docx_elements(docx_path)
    paragraph_text = "\n".join(e.text for e in elements if e.element_type != "table")
    tables = [e.metadata for e in elements if e.element_type == "table"]
    return {"paragraph_text": paragraph_text, "tables": tables}


def _save_parent_index(parents: list[ParentChunk], skipped_files: list[Path], use_llm: bool) -> None:
    RUNTIME_KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    payload = {
        "version": 2,
        "build_at": date.today().isoformat(),
        "parents": [p.to_dict() for p in parents],
        "skipped_files": [str(p) for p in skipped_files],
        "models": {
            "chunk_model": CHUNK_MODEL,
            "embed_model": EMBED_MODEL,
            "vision_model": VISION_MODEL,
            "use_llm_chunking": str(use_llm),
        },
    }
    PARENT_INDEX_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def _load_parent_index() -> dict[str, ParentChunk]:
    if not PARENT_INDEX_FILE.exists():
        return {}
    try:
        payload = json.loads(PARENT_INDEX_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}
    parents = payload.get("parents", [])
    out: dict[str, ParentChunk] = {}
    for item in parents:
        try:
            parent = ParentChunk.from_dict(item)
        except Exception:
            continue
        out[parent.parent_id] = parent
    return out


def _build_chunks_from_files(files: list[Path], use_llm: bool = True) -> tuple[list[ParentChunk], list[ChildChunk], dict[str, int]]:
    parents: list[ParentChunk] = []
    children: list[ChildChunk] = []
    stats = {"documents": 0, "parents": 0, "children": 0, "tables": 0, "errors": 0}
    for source in files:
        try:
            elements = _load_document_elements(source)
        except Exception as exc:
            stats["errors"] += 1
            print(f"[RAG] 跳过解析失败文件: {source.name} ({exc})", file=sys.stderr)
            continue
        if not elements:
            continue
        stats["documents"] += 1
        doc_parents = _elements_to_parents(elements)
        parents.extend(doc_parents)
        for parent in doc_parents:
            if parent.chunk_type == "table":
                stats["tables"] += 1
            children.extend(_children_from_parent(parent, use_llm=use_llm))
    stats["parents"] = len(parents)
    stats["children"] = len(children)
    return parents, children, stats


def build_knowledge_base(docx_path: str | Path | None = None, use_llm: bool | None = None) -> dict[str, Any]:
    """Build the parent/child knowledge base.

    Passing a file path keeps the old single-file behavior. Passing None scans
    the repository-level Chinese knowledge directory.
    """
    supported_files, skipped_files = _discover_knowledge_files(docx_path)
    if not supported_files:
        target = Path(docx_path) if docx_path else SOURCE_KNOWLEDGE_DIR
        return {"status": "error", "chunks": 0, "message": f"未找到可解析知识库文件: {target}"}

    if use_llm is None:
        use_llm = USE_LLM_CHUNKING
    mode = "LLM 语义边界" if use_llm else "规则语义边界"
    print(f"[RAG] 发现 {len(supported_files)} 个知识库文件，跳过 {len(skipped_files)} 个旧格式文件，分割模式：{mode}", file=sys.stderr)
    parents, children, stats = _build_chunks_from_files(supported_files, use_llm=use_llm)
    if not children:
        return {"status": "error", "chunks": 0, "message": "知识库分块为空"}

    RUNTIME_KNOWLEDGE_DIR.mkdir(parents=True, exist_ok=True)
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(str(CHROMA_DIR))
    try:
        client.delete_collection("party_rules")
    except Exception:
        pass
    collection = client.create_collection("party_rules")

    build_at = date.today().isoformat()
    documents = [child.content for child in children]
    embedding_documents = [child.embedding_text for child in children]
    ids = [child.child_id for child in children]
    metadatas = [child.to_metadata(i, build_at) for i, child in enumerate(children)]

    print(f"[RAG] 构建向量索引：{len(children)} 个 child chunks", file=sys.stderr)
    batch_size = 16
    for start in range(0, len(documents), batch_size):
        end = min(start + batch_size, len(documents))
        batch_docs = documents[start:end]
        batch_embedding_docs = embedding_documents[start:end]
        batch_meta = metadatas[start:end]
        batch_ids = ids[start:end]
        embeddings = _ollama_embed(batch_embedding_docs)
        if embeddings and len(embeddings) == len(batch_docs):
            collection.add(documents=batch_docs, metadatas=batch_meta, ids=batch_ids, embeddings=embeddings)
        else:
            print(f"[WARN] Ollama 嵌入失败，批次 {start // batch_size + 1} 使用 Chroma 默认嵌入回退", file=sys.stderr)
            collection.add(documents=batch_docs, metadatas=batch_meta, ids=batch_ids)
        progress = min(100, int(end / len(documents) * 100))
        print(f"  [RAG] 向量化进度: {progress}% ({end}/{len(documents)})", file=sys.stderr)

    CHUNK_INDEX_FILE.write_text(
        json.dumps({"version": 2, "documents": documents, "metadatas": metadatas, "ids": ids}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    _save_parent_index(parents, skipped_files, use_llm)

    skipped_note = f"，跳过旧 .doc 文件 {len(skipped_files)} 个" if skipped_files else ""
    message = (
        f"知识库构建完成：文档 {stats['documents']} 个，parent {len(parents)} 个，"
        f"child {len(children)} 个，表格 parent {stats['tables']} 个{skipped_note}"
    )
    print(f"[RAG] {message}", file=sys.stderr)
    return {
        "status": "ok",
        "chunks": len(children),
        "parents": len(parents),
        "documents": stats["documents"],
        "table_parents": stats["tables"],
        "skipped": len(skipped_files),
        "message": message,
    }


_ce_model = None


def _get_ce_model():
    global _ce_model
    if _ce_model is None:
        from sentence_transformers import CrossEncoder

        _ce_model = CrossEncoder(RERANK_MODEL_NAME)
    return _ce_model


_RERANK_SYSTEM_PROMPT = """你是党务知识检索的精确排序专家。
请只输出每行一个 `段落ID: 分数`，分数 0-10，按内容是否能回答用户查询评分。
"""


def _llm_rerank(chunks: list[dict[str, Any]], query: str, top_k: int = 5) -> list[dict[str, Any]]:
    if not chunks:
        return []
    candidates = []
    for i, chunk in enumerate(chunks, 1):
        candidates.append(f"段落{i}: {chunk['title']}\n{chunk['content'][:700]}")
    prompt = f"用户查询：{query}\n\n" + "\n\n".join(candidates)
    result = _call_llm(_RERANK_SYSTEM_PROMPT, prompt, max_tokens=512, model=MODEL)
    scores: dict[int, float] = {}
    for line in result.splitlines():
        match = re.search(r"段落\s*(\d+)\s*[:：]\s*([0-9.]+)", line)
        if match:
            scores[int(match.group(1)) - 1] = float(match.group(2))
    for i, chunk in enumerate(chunks):
        score = scores.get(i, chunk.get("score", 0.0) * 10)
        chunk["rerank_score"] = score
        chunk["score"] = round(score / 10, 4)
        chunk["source"] = "hybrid+rerank(llm)"
    return sorted(chunks, key=lambda x: x.get("rerank_score", 0), reverse=True)[:top_k]


def _ce_rerank(chunks: list[dict[str, Any]], query: str, top_k: int = 5) -> list[dict[str, Any]]:
    if not chunks:
        return []
    try:
        model = _get_ce_model()
        pairs = [(query, f"{c['title']}\n{c['content'][:700]}") for c in chunks]
        scores = model.predict(pairs)
        for i, chunk in enumerate(chunks):
            score = float(scores[i])
            chunk["rerank_score"] = score
            chunk["score"] = round(score, 4)
            chunk["source"] = "hybrid+rerank(ce)"
        return sorted(chunks, key=lambda x: x["rerank_score"], reverse=True)[:top_k]
    except Exception as exc:
        print(f"[WARN] Cross-Encoder 精排失败 ({exc})，回退到 LLM/融合排序", file=sys.stderr)
        return _llm_rerank(chunks, query, top_k=top_k)


class VectorRAG:
    def __init__(self) -> None:
        self._chroma_collection = None
        self._bm25: BM25Okapi | None = None
        self._bm25_docs: list[dict[str, Any]] = []
        self._parents: dict[str, ParentChunk] = {}
        self._initialized = False

    def _load_bm25(self) -> None:
        self._bm25 = None
        self._bm25_docs = []
        if not CHUNK_INDEX_FILE.exists():
            return
        data = json.loads(CHUNK_INDEX_FILE.read_text(encoding="utf-8"))
        documents = data.get("documents", [])
        metadatas = data.get("metadatas", [])
        ids = data.get("ids", [])
        tokenized = [_tokenize(doc) for doc in documents]
        self._bm25 = BM25Okapi(tokenized) if tokenized else None
        self._bm25_docs = [
            {"content": doc, "metadata": meta, "id": doc_id}
            for doc, meta, doc_id in zip(documents, metadatas, ids)
        ]

    def _load_parents(self) -> None:
        self._parents = _load_parent_index()

    def _ensure_initialized(self) -> bool:
        if self._initialized:
            return True
        if not CHROMA_DIR.exists():
            return False
        try:
            client = chromadb.PersistentClient(str(CHROMA_DIR))
            self._chroma_collection = client.get_collection("party_rules")
            self._load_bm25()
            self._load_parents()
            self._initialized = True
            return True
        except Exception:
            return False

    def _with_parent_context(self, item: dict[str, Any], metadata: dict[str, Any] | None = None) -> dict[str, Any]:
        meta = metadata or item.get("metadata") or {}
        parent_id = meta.get("parent_id") or item.get("parent_id", "")
        parent = self._parents.get(parent_id)
        context = ""
        if parent:
            context = parent.content[:900].strip()
            if len(parent.content) > 900:
                context += "..."
        citation = meta.get("citation") or item.get("citation") or item.get("title", "")
        item.update(
            {
                "parent_id": parent_id,
                "doc_name": meta.get("doc_name", ""),
                "heading_path": meta.get("heading_path", ""),
                "article_no": meta.get("article_no", ""),
                "chunk_type": meta.get("chunk_type", "text"),
                "citation": citation,
                "parent_context": context,
                "metadata": meta,
            }
        )
        return item

    def search(
        self,
        query: str,
        top_k: int = DEFAULT_TOP_K,
        alpha: float = HYBRID_ALPHA,
        rerank: bool = DEFAULT_RERANK,
    ) -> list[dict[str, Any]]:
        if not self._ensure_initialized():
            return []

        candidate_k = HYBRID_TOP_K if rerank else top_k
        dense_results: dict[str, dict[str, Any]] = {}
        try:
            query_embedding = _ollama_embed([query])
            if query_embedding:
                resp = self._chroma_collection.query(
                    query_embeddings=query_embedding,
                    n_results=candidate_k + 5,
                    include=["documents", "metadatas", "distances"],
                )
            else:
                resp = self._chroma_collection.query(
                    query_texts=[query],
                    n_results=candidate_k + 5,
                    include=["documents", "metadatas", "distances"],
                )
            for i, doc_id in enumerate(resp.get("ids", [[]])[0]):
                meta = resp["metadatas"][0][i]
                dist = resp["distances"][0][i] if resp.get("distances") else 1.0
                sim = max(0.0, 1.0 - float(dist) / 2.0)
                dense_results[doc_id] = self._with_parent_context(
                    {
                        "id": doc_id,
                        "title": meta.get("title", "未知"),
                        "content": resp["documents"][0][i],
                        "dense_score": round(sim, 4),
                    },
                    meta,
                )
        except Exception:
            pass

        sparse_results: dict[str, dict[str, Any]] = {}
        if self._bm25 and self._bm25_docs:
            tokens = _tokenize(query)
            if tokens:
                scores = self._bm25.get_scores(tokens)
                max_score = max(scores) if len(scores) and max(scores) > 0 else 1.0
                ranked = sorted(enumerate(scores), key=lambda x: x[1], reverse=True)[: candidate_k + 5]
                for i, score in ranked:
                    if score <= 0:
                        continue
                    doc = self._bm25_docs[i]
                    meta = doc["metadata"]
                    sparse_results[doc["id"]] = self._with_parent_context(
                        {
                            "id": doc["id"],
                            "title": meta.get("title", "未知"),
                            "content": doc["content"],
                            "sparse_score": round(float(score) / float(max_score), 4),
                        },
                        meta,
                    )

        fused: list[dict[str, Any]] = []
        all_ids = list(set(dense_results) | set(sparse_results))

        if FUSION_MODE == "rrf":
            # ── RRF (Reciprocal Rank Fusion) ──
            dense_ranked = sorted(
                dense_results.items(), key=lambda x: x[1].get("dense_score", 0), reverse=True
            )
            sparse_ranked = sorted(
                sparse_results.items(), key=lambda x: x[1].get("sparse_score", 0), reverse=True
            )
            dense_ranks = {doc_id: idx for idx, (doc_id, _) in enumerate(dense_ranked)}
            sparse_ranks = {doc_id: idx for idx, (doc_id, _) in enumerate(sparse_ranked)}

            for doc_id in all_ids:
                d = dense_results.get(doc_id, {})
                s = sparse_results.get(doc_id, {})
                base = d or s
                dr = dense_ranks.get(doc_id, len(dense_ranked))
                sr = sparse_ranks.get(doc_id, len(sparse_ranked))
                rrf_score = 1.0 / (RRF_K + dr) + 1.0 / (RRF_K + sr)
                fused.append({
                    **base,
                    "id": doc_id,
                    "title": base.get("title", "未知"),
                    "content": base.get("content", ""),
                    "score": round(rrf_score, 4),
                    "dense_score": d.get("dense_score", 0.0),
                    "sparse_score": s.get("sparse_score", 0.0),
                    "source": "hybrid+rrf",
                })
        else:
            # ── 原始线性加权融合 ──
            for doc_id in all_ids:
                d = dense_results.get(doc_id, {})
                s = sparse_results.get(doc_id, {})
                base = d or s
                dense_score = d.get("dense_score", 0.0)
                sparse_score = s.get("sparse_score", 0.0)
                hybrid_score = alpha * dense_score + (1 - alpha) * sparse_score
                fused.append({
                    **base,
                    "id": doc_id,
                    "title": base.get("title", "未知"),
                    "content": base.get("content", ""),
                    "score": round(hybrid_score, 4),
                    "dense_score": dense_score,
                    "sparse_score": sparse_score,
                    "source": "hybrid+linear",
                })
        fused.sort(key=lambda x: x["score"], reverse=True)

        if rerank and fused:
            candidates = fused[:HYBRID_TOP_K]
            if RERANK_MODE == "cross-encoder":
                return _ce_rerank(candidates, query, top_k=top_k)
            if RERANK_MODE == "llm":
                return _llm_rerank(candidates, query, top_k=top_k)
        return fused[:top_k]

    def get_relevant_procedure(self, step_name: str, top_k: int = 2) -> list[dict[str, Any]]:
        results = self.search(f"贵州省发展党员工作规程 {step_name}", top_k=top_k, rerank=True)
        if len(results) < top_k:
            seen = {r["id"] for r in results}
            for item in self.search(step_name, top_k=top_k * 2, rerank=True):
                if item["id"] not in seen:
                    results.append(item)
                    seen.add(item["id"])
                if len(results) >= top_k:
                    break
        return results[:top_k]

    def format_citations(self, results: list[dict[str, Any]], min_score: float = 0.15) -> str:
        valid = [r for r in results if r.get("score", 0) >= min_score]
        if not valid:
            return ""
        lines = ["知识库引用："]
        for i, result in enumerate(valid, 1):
            citation = result.get("citation") or result.get("title", "未知")
            lines.append(f"\n{i}. {citation} (相关度: {result.get('score', 0):.0%})")
            evidence = result.get("content", "")[:180].strip()
            if len(result.get("content", "")) > 180:
                evidence += "..."
            lines.append(f"   证据：{evidence}")
        lines.append(f"\n--- 共检索 {len(valid)} 条相关规程 ---")
        return "\n".join(lines)

    def enhanced_system_context(self, query: str, top_k: int = 3) -> str:
        results = self.search(query, top_k=top_k, rerank=True)
        valid = [r for r in results if r.get("score", 0) >= 0.15]
        knowledge = []
        for i, result in enumerate(valid, 1):
            content = result.get("content", "").strip()
            context = result.get("parent_context", "").strip()
            knowledge.append(
                {
                    "id": str(i),
                    "title": result.get("title", ""),
                    "content": content[:300] + ("..." if len(content) > 300 else ""),
                    "evidence": content,
                    "context": context,
                    "citation": result.get("citation", result.get("title", "")),
                    "chunk_type": result.get("chunk_type", "text"),
                }
            )
        return json.dumps({"knowledge": knowledge}, ensure_ascii=False)


def rebuild_command() -> None:
    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
    if CHUNK_INDEX_FILE.exists():
        CHUNK_INDEX_FILE.unlink()
    if PARENT_INDEX_FILE.exists():
        PARENT_INDEX_FILE.unlink()
    result = build_knowledge_base()
    print(f"[RAG] {result['message']}", file=sys.stderr)


def status_command() -> dict[str, Any]:
    rag = VectorRAG()
    if not rag._ensure_initialized():
        return {"ready": False, "chunks": 0, "parents": 0, "message": "知识库未构建，请先运行 uv run python -m partymate.tools.rag --rebuild"}
    try:
        chunk_count = rag._chroma_collection.count()
        parent_count = len(rag._parents)
        return {
            "ready": True,
            "chunks": chunk_count,
            "parents": parent_count,
            "message": f"知识库就绪，{chunk_count} 个 child chunks，{parent_count} 个 parent chunks",
        }
    except Exception as exc:
        return {"ready": False, "chunks": 0, "parents": 0, "message": str(exc)}


_rag_instance: VectorRAG | None = None


def _get_rag() -> VectorRAG:
    global _rag_instance
    if _rag_instance is None:
        _rag_instance = VectorRAG()
    return _rag_instance


def search(query: str, top_k: int = DEFAULT_TOP_K) -> list[dict[str, Any]]:
    return _get_rag().search(query, top_k=top_k)


def format_citations(results: list[dict[str, Any]], min_score: float = 0.15) -> str:
    return _get_rag().format_citations(results, min_score=min_score)


def search_with_fallback(query: str, top_k: int = 3) -> list[dict[str, Any]]:
    return _get_rag().search(query, top_k=top_k)


def get_rag_status() -> dict[str, Any]:
    return status_command()


if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] == "--rebuild":
        rebuild_command()
    elif len(sys.argv) > 2 and sys.argv[1] == "--search":
        rag = _get_rag()
        results = rag.search(sys.argv[2])
        print(rag.format_citations(results), file=sys.stderr)
    elif len(sys.argv) > 2 and sys.argv[1] == "--procedure":
        rag = _get_rag()
        results = rag.get_relevant_procedure(sys.argv[2])
        print(rag.format_citations(results), file=sys.stderr)
    else:
        print(
            "用法:\n"
            "  uv run python -m partymate.tools.rag --rebuild\n"
            "  uv run python -m partymate.tools.rag --search <query>\n"
            "  uv run python -m partymate.tools.rag --procedure <step>",
            file=sys.stderr,
        )
